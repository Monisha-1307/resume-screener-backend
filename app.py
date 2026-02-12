import os
import urllib
from flask import Flask, request, jsonify
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from flask_jwt_extended import JWTManager
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pdfplumber
from docx import Document
import traceback

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "https://resume-screener-frontend-1.onrender.com"}})

# -------------------------------
# Environment variables
# -------------------------------
app.config['JWT_SECRET_KEY'] = os.environ.get('JWT_SECRET_KEY', 'fallbacksecret')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///default.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
jwt = JWTManager(app)

# -------------------------------
# Database models
# -------------------------------
class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200), nullable=False)
    content = db.Column(db.Text, nullable=False)

class Job(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, nullable=False)

class Comparison(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    resume_id = db.Column(db.Integer, db.ForeignKey('resume.id'), nullable=False)
    job_id = db.Column(db.Integer, db.ForeignKey('job.id'), nullable=False)
    score = db.Column(db.Float, nullable=False)
    keywords = db.Column(db.Text, nullable=True)

# -------------------------------
# Health check
# -------------------------------
@app.route('/ping')
def ping():
    return jsonify({"status": "ok"})

@app.route('/')
def home():
    return jsonify({"message": "Resume Screener Backend is running!"})

# -------------------------------
# Utility function
# -------------------------------
def calculate_similarity_with_keywords(resume_text, job_text):
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform([resume_text, job_text])
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]

    resume_words = set(resume_text.lower().split())
    job_words = set(job_text.lower().split())
    common_words = list(resume_words.intersection(job_words))

    return round(similarity * 100, 2), common_words

# -------------------------------
# Upload resume
# -------------------------------
@app.route('/upload_resume', methods=['POST'])
def upload_resume():
    print("Received upload_resume request")
    if 'resume' not in request.files:
        print("No resume file in request")
        return jsonify({"error": "No resume file uploaded"}), 400

    file = request.files['resume']
    filename = file.filename.lower()
    text = ""

    try:
        if filename.endswith(".pdf"):
            file.seek(0)
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

        elif filename.endswith(".docx"):
            file.seek(0)
            doc = Document(file)  # ✅ direct file object
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"

        else:
            file.seek(0)
            text = file.read().decode('utf-8', errors='ignore')

        if not text.strip():
            print("No text extracted from resume")
            return jsonify({"error": "No text extracted from resume"}), 500

        new_resume = Resume(filename=filename, content=text)
        db.session.add(new_resume)
        db.session.commit()

        print(f"Stored resume {filename} with id={new_resume.id}, length={len(text)}")
        return jsonify({"resume_text": text, "resume_id": new_resume.id})

    except Exception as e:
        print("Extraction failed:", str(e))
        traceback.print_exc()
        return jsonify({"error": f"Failed to extract text: {str(e)}"}), 500

# -------------------------------
# Other routes (jobs, match, summary, etc.)
# -------------------------------
# ... keep the rest of your code unchanged ...
# (add_job, match, match_multiple, resume_summary, get_resumes, get_jobs, get_comparisons, list_routes)

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)
